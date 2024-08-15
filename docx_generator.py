import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import  WD_SECTION_START
import streamlit as st
import docx
import pandas as pd
import io
import time
import io
import urllib.request
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image,UnidentifiedImageError

# your avatar URL as example
url = ('https://drive.google.com/uc?export=download&id=1nB6n0ENZPHquN6hYl3sf7dlXGn5aP7kD')
try:
    content = urllib.request.urlopen(url).read()
    cover_image = Image.open(io.BytesIO(content))
    cover_image = cover_image.convert("RGB")
except UnidentifiedImageError:
    st.error("The image could not be identified or is not in a valid format.")
    cover_image = None
except Exception as e:
    st.error(f"An error occurred while fetching the image: {e}")
    cover_image = None

def generate_docx_report(df_long, indicators):
    # Dictionary to store images by indicator number
    images_by_goal = {}
    message_placeholder = st.empty()
    message_placeholder.success("Report generation started. Please wait for a minute to generate the report.")
    time.sleep(5)
    message_placeholder.empty()

    for indicator in indicators:
        print(f"Processing indicator: {indicator}")
        indicator_df = df_long[df_long['Indicator'] == indicator]
        if indicator_df.empty:
            continue

        indicator_number = indicator_df['Indicator Number'].iloc[0]
        goal_index = int(indicator_number.split('.')[0])  # Extract goal number from indicator number

        # Plotting
        plt.figure(figsize=(20, 10))

        for sub_category in indicator_df['Sub Category'].unique():
            sub_category_df = indicator_df[indicator_df['Sub Category'] == sub_category]
            if not sub_category_df.empty:
                plt.plot(sub_category_df['Year'], sub_category_df['Value'], marker='o', label=sub_category)
            else:
                if str(indicator_df['Target Value']) != 'nan':
                    plt.plot(indicator_df['Year'], indicator_df['Value'], marker='o')

        plt.plot(indicator_df['Year'], indicator_df['Target Value'], linestyle='--', label='Target Value',
                 color='purple')

        plt.title(f"{indicator_number}: {indicator}")
        plt.xlabel('Year')
        plt.ylabel('Value')
        plt.legend()

        # Save plot to buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        plt.close()
        img_buffer.seek(0)

        if goal_index not in images_by_goal:
            images_by_goal[goal_index] = []
        images_by_goal[goal_index].append(img_buffer)

    message_placeholder.success("Processing almost done. The report will be generated in less than a minute.")
    time.sleep(5)
    message_placeholder.empty()

    # Create the DOCX report
    docx_file = create_docx_report(images_by_goal,cover_image)

    return docx_file

def add_border(section, color):
    """Add border color for a section based on the SDG goal."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '30')  # Border thickness in half-points (10px = 7.5 points)
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), color)
        pgBorders.append(border)
    sectPr.append(pgBorders)

def create_docx_report(images_by_goal, cover_image):
    doc = Document()

    if cover_image:
        cover_buffer = io.BytesIO()
        cover_image.save(cover_buffer, format='PNG')
        cover_buffer.seek(0)
        doc.add_picture(cover_buffer, width=Inches(6), height=Inches(8))
        doc.add_page_break()

    sdg_goals = [
        "No Poverty", "Zero Hunger", "Good Health and Well-being", "Quality Education",
        "Gender Equality", "Clean Water and Sanitation", "Affordable and Clean Energy",
        "Decent Work and Economic Growth", "Industry, Innovation, and Infrastructure",
        "Reduced Inequality", "Sustainable Cities and Communities",
        "Responsible Consumption and Production", "Climate Action", "Life Below Water",
        "Life on Land", "Peace, Justice, and Strong Institutions", "Partnerships for the Goals"
    ]

    # Create index for SDG goals
    index_heading = doc.add_heading('Sustainable Development Goals', level=1)
    index_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create a table for SDG goals
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S.No'
    hdr_cells[1].text = 'SDG Goal'
    hdr_cells[2].text = 'Page No.'

    page_number = 3
    for i, goal in enumerate(sdg_goals, start=1):
        l = len(images_by_goal.get(i, []))
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = goal
        row_cells[2].text = str(page_number)
        page_number += l // 2 + 1 if l % 2 else l // 2

    # Add a section break to start the new section after the index, but don't insert a page break
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)

    # Unlink footer and set up page numbering to start from 1
    section.footer.is_linked_to_previous = False  # Unlink from previous section
    section.start_type = WD_SECTION_START.CONTINUOUS
    section.page_number_start = 1  # Start numbering from 1

    # Add page number to the footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adding page number using an OxmlElement
    run = footer_paragraph.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

    sdg_colors = [
        "E5243B", "DDA63A", "4C9F38", "C5192D", "FF3A21", "26BDE2", "FCC30B",
        "A21942", "FD6925", "DD1367", "FD9D24", "BF8B2E", "3F7E44", "0A97D9",
        "56C02B", "00689D", "19486A"
    ]

    # Iterate through SDG goals and their corresponding images
    for goal_index, goal_images in images_by_goal.items():
        section = doc.add_section()
        add_border(section, sdg_colors[goal_index - 1])
        # doc.add_page_break()

        # Add goal heading
        goal_heading = doc.add_heading(f'SDG {goal_index}: {sdg_goals[goal_index - 1]}', level=2)
        goal_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for img_buffer in goal_images:
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6), height=Inches(4))

    # Use buffer for the DOCX file
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer
